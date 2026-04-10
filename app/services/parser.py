import re
import pdfplumber
from docx import Document
from docx.oxml.ns import qn  # NEW IMPORT REQUIRED FOR COLOR EXTRACTION
# ==========================================
# SHARED UTILITIES
# ==========================================

def clean_text(text):
    """Cleans up newlines and excessive whitespace from cell text."""
    if not text: return ""
    return re.sub(r'\s+', ' ', str(text)).strip()

# ==========================================
# DOCX PARSER 
# ==========================================

# --- NEW HELPER FUNCTION ---
def is_border_removed(cell, side):
    """Checks the XML of the cell to see if a specific border was explicitly removed."""
    borders = cell._tc.xpath(f'.//w:tcBorders/w:{side}')
    if borders:
        val = borders[0].get(qn('w:val'))
        if val:
            return val.lower() in ['nil', 'none', 'clear', '']
    return False

def process_docx(file_path):
    doc = Document(file_path)
    if not doc.tables: return []
    
    biggest_table = max(doc.tables, key=lambda t: len(t.rows) * len(t.columns))
    
    # 1. Map the grid with Color and Border States
    grid = []
    for row in biggest_table.rows:
        grid_row = []
        for cell in row.cells:
            color = None
            shading_elms = cell._tc.xpath('.//w:shd')
            if shading_elms:
                fill = shading_elms[0].get(qn('w:fill'))
                if fill and fill.upper() not in ['AUTO', 'FFFFFF', '000000', 'CLEAR']:
                    color = fill

            grid_row.append({
                'text': clean_text(cell.text),
                'id': cell._element,
                'color': color,
                'bottom_removed': is_border_removed(cell, 'bottom'),
                'top_removed': is_border_removed(cell, 'top')
            })
        grid.append(grid_row)
    
    # 2. Find Headers
    header_idx = -1
    for i, row in enumerate(grid):
        if any("MONDAY" in str(cell['text']).upper() for cell in row):
            header_idx = i
            break
            
    if header_idx == -1: return []
    
    headers = [str(c['text']).strip().upper() for c in grid[header_idx]]
    day_columns = {h: idx for idx, h in enumerate(headers) if h in ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY", "SATURDAY"]}

    completed_blocks = []
    active_blocks = {day: None for day in day_columns.keys()}

    last_valid_start = None
    last_valid_end = None

    # 3. Process Rows
    for row_idx in range(header_idx + 1, len(grid)):
        row = grid[row_idx]
        if not row: continue
        
        time_col = str(row[0]['text']).upper()
        if "SUMMARY" in time_col or "REGULAR TEACHING" in time_col: break

        matches = re.findall(r'(\d{1,2})[^\dA-Za-z]*(\d{2})\s*(AM|PM)', str(time_col), re.IGNORECASE)
        times = []
        for h, m, p in matches:
            hour = int(h)
            if hour == 12: hour = 12
            elif p.upper() == 'PM' and hour < 12: hour += 12
            times.append(f"{hour:02d}:{m}")
            
        start_time = times[0] if len(times) >= 1 else None
        end_time = times[-1] if len(times) >= 2 else start_time
        
        if start_time and end_time:
            last_valid_start = start_time
            last_valid_end = end_time
        else:
            start_time = last_valid_start
            end_time = last_valid_end
            
        if not start_time or not end_time: continue

        for day, col_idx in day_columns.items():
            if col_idx >= len(row): continue
            
            cell_data = row[col_idx]
            cell_text = cell_data['text']
            cell_id = cell_data['id']
            cell_color = cell_data['color']
            top_removed = cell_data['top_removed']
            bottom_removed = cell_data['bottom_removed']
            
            is_empty = not cell_text or cell_text.lower() in ['break', 'lunch', 'n/a', '', '`']
            has_fill = cell_color is not None

            same_block = False
            if active_blocks[day]:
                if active_blocks[day]['id'] == cell_id:
                    same_block = True
                elif has_fill and active_blocks[day]['color'] == cell_color:
                    if active_blocks[day]['bottom_removed'] or top_removed:
                        same_block = True

            # KEY FIX: A colored block can start even if the cell is currently empty
            if not is_empty or same_block or has_fill:
                if same_block:
                    active_blocks[day]['end'] = end_time
                    active_blocks[day]['bottom_removed'] = bottom_removed 
                    
                    if cell_text and cell_text not in active_blocks[day]['event']:
                        if active_blocks[day]['event']:
                            active_blocks[day]['event'] += f" {cell_text}"
                        else:
                            active_blocks[day]['event'] = cell_text
                else:
                    if active_blocks[day]: 
                        completed_blocks.append(active_blocks[day])
                    
                    active_blocks[day] = {
                        'day': day, 
                        'event': cell_text if not is_empty else "", 
                        'start': start_time, 
                        'end': end_time,
                        'id': cell_id,
                        'color': cell_color,
                        'bottom_removed': bottom_removed
                    }
            else:
                if active_blocks[day]:
                    completed_blocks.append(active_blocks[day])
                    active_blocks[day] = None

    # 4. Cleanup
    for block in active_blocks.values():
        if block:
            completed_blocks.append(block)
            
    final_blocks = []
    for block in completed_blocks:
        # Filter out blocks that were colored but NEVER contained text
        if block['event'].strip():
            if 'id' in block: del block['id']
            if 'color' in block: del block['color']
            if 'bottom_removed' in block: del block['bottom_removed']
            final_blocks.append(block)

    return final_blocks


def process_pdf(file_path):
    completed_blocks = []
    
    with pdfplumber.open(file_path) as pdf:
        page = pdf.pages[0]
        words = page.extract_words()
        
        headers = ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY", "SATURDAY"]
        header_words = [w for w in words if w['text'].upper() in headers]
        if not header_words: return []
        
        # 1. Map Columns using Midpoints
        col_cx = {w['text'].upper(): (w['x0'] + w['x1']) / 2 for w in header_words}
        sorted_days = sorted(col_cx.keys(), key=lambda d: col_cx[d])
        monday_x0 = next(w['x0'] for w in header_words if w['text'].upper() == "MONDAY")
        
        col_bounds = {}
        for i, day in enumerate(sorted_days):
            start_x = monday_x0 - 15 if i == 0 else (col_cx[sorted_days[i-1]] + col_cx[day]) / 2
            end_x = (col_cx[day] + col_cx[sorted_days[i+1]]) / 2 if i+1 < len(sorted_days) else page.width
            col_bounds[day] = (start_x, end_x)
            
        # 2. Strict Boundary for the Main Grid
        table_top = min(w['top'] for w in header_words) - 5
        summary_words = [w for w in words if "SUMMARY" in w['text'].upper() or "TEACHING" in w['text'].upper()]
        table_bottom = min(w['top'] for w in summary_words) - 5 if summary_words else page.height

        table_words = [w for w in words if w['top'] >= table_top and w['bottom'] <= table_bottom]

        # 3. Y-axis Dynamic Time Grid (Reads actual left-column times)
        time_col_words = [w for w in table_words if w['x1'] < monday_x0]
        time_col_words.sort(key=lambda w: w['top'])
        
        lines = []
        current_line = []
        for w in time_col_words:
            if not current_line:
                current_line.append(w)
            elif abs(w['top'] - current_line[0]['top']) < 10:
                current_line.append(w)
            else:
                lines.append(current_line)
                current_line = [w]
        if current_line: lines.append(current_line)
        
        time_grid = []
        for line in lines:
            line.sort(key=lambda w: w['x0'])
            text = " ".join([w['text'] for w in line])
            matches = re.findall(r'(\d{1,2}:\d{2})\s*(AM|PM)', text, re.IGNORECASE)
            
            if len(matches) >= 2:
                def to_24h(t, p):
                    h, m = map(int, t.split(':'))
                    if h == 12: h = 12 if p.upper() == 'PM' else 0
                    elif p.upper() == 'PM' and h < 12: h += 12
                    return f"{h:02d}:{m:02d}"
                    
                start_time = to_24h(matches[0][0], matches[0][1])
                end_time = to_24h(matches[-1][0], matches[-1][1])
                
                top_y = min(w['top'] for w in line)
                bottom_y = max(w['bottom'] for w in line)
                
                time_grid.append({
                    'start': start_time,
                    'end': end_time,
                    'top': top_y,
                    'bottom': bottom_y,
                    'center': (top_y + bottom_y) / 2
                })

        if not time_grid: return []

        def get_time_from_y(y, is_start=True):
            if is_start:
                closest_slot = min(time_grid, key=lambda slot: abs(slot['top'] - y))
                return closest_slot['start']
            else:
                closest_slot = min(time_grid, key=lambda slot: abs(slot['bottom'] - y))
                return closest_slot['end']

        grid_top = time_grid[0]['top']
        grid_bottom = time_grid[-1]['bottom']

        # 4. Extract colored background blocks & Remove duplicates
        raw_rects = [r for r in page.rects if r['top'] >= grid_top - 10 and r['bottom'] <= grid_bottom + 10 and r['height'] > 5 and 20 < r['width'] < (page.width * 0.4)]
        
        unique_rects = []
        for r in raw_rects:
            r_center_x = (r['x0'] + r['x1']) / 2
            r_center_y = (r['top'] + r['bottom']) / 2
            is_dup = False
            for idx, ur in enumerate(unique_rects):
                ur_center_x = (ur['x0'] + ur['x1']) / 2
                ur_center_y = (ur['top'] + ur['bottom']) / 2
                if abs(r_center_x - ur_center_x) < 5 and abs(r_center_y - ur_center_y) < 5:
                    is_dup = True
                    if (r['width'] * r['height']) > (ur['width'] * ur['height']):
                        unique_rects[idx] = r
                    break
            if not is_dup:
                unique_rects.append(r)

        # 5. Assign every word a tag mapping it to its specific Colored Box
        def get_rect_idx(w):
            w_cx = (w['x0'] + w['x1']) / 2
            w_cy = (w['top'] + w['bottom']) / 2
            for idx, r in enumerate(unique_rects):
                if r['x0'] - 5 <= w_cx <= r['x1'] + 5 and r['top'] - 5 <= w_cy <= r['bottom'] + 5:
                    return idx
            return -1

        for w in table_words:
            w['rect_idx'] = get_rect_idx(w)

        # 6. Read data strictly within columns
        for day, (x0, x1) in col_bounds.items():
            col_words = []
            for w in table_words:
                w_cx = (w['x0'] + w['x1']) / 2
                if x0 <= w_cx < x1 and w['top'] > grid_top - 10:
                    text_upper = clean_text(w['text']).upper()
                    if text_upper not in headers and text_upper not in ['BREAK', 'LUNCH', 'N/A', '']:
                        col_words.append(w)
                        
            col_words.sort(key=lambda w: w['top'])
            
            # Cluster floating text strictly restricted by colored boxes
            clusters = []
            current_cluster = []
            for w in col_words:
                if not current_cluster:
                    current_cluster.append(w)
                else:
                    last_w = max(current_cluster, key=lambda x: x['bottom'])
                    
                    same_rect = (w['rect_idx'] == last_w['rect_idx'])
                    
                    if same_rect and w['rect_idx'] != -1 and w['top'] - last_w['bottom'] <= 35:
                        current_cluster.append(w)
                    elif w['rect_idx'] == -1 and last_w['rect_idx'] == -1 and w['top'] - last_w['bottom'] <= 15:
                        current_cluster.append(w)
                    else:
                        clusters.append(current_cluster)
                        current_cluster = [w]
                        
            if current_cluster: clusters.append(current_cluster)

            parsed_clusters = []
            for c in clusters:
                c.sort(key=lambda w: (w['top'], w['x0']))
                text = clean_text(" ".join([w['text'] for w in c]))
                if not text: continue
                
                parsed_clusters.append({
                    'text': text,
                    'top': min(w['top'] for w in c),
                    'bottom': max(w['bottom'] for w in c),
                    'rect_idx': c[0]['rect_idx']
                })

            for cluster in parsed_clusters:
                matched_rect = None
                if cluster['rect_idx'] != -1:
                    matched_rect = unique_rects[cluster['rect_idx']]

                if matched_rect:
                    # Use the boundaries of the colored box
                    start_time = get_time_from_y(matched_rect['top'], is_start=True)
                    end_time = get_time_from_y(matched_rect['bottom'], is_start=False)
                else:
                    # Use the text's own bounding box
                    start_time = get_time_from_y(cluster['top'], is_start=True)
                    end_time = get_time_from_y(cluster['bottom'], is_start=False)
                
                if start_time and end_time and start_time != end_time:
                    completed_blocks.append({
                        'day': day,
                        'event': cluster['text'],
                        'start': start_time,
                        'end': end_time
                    })
                    
    return completed_blocks

# ==========================================
# MAIN EXECUTION testing
# ==========================================

def run():
    files = ["Engr. Carlo.pdf"]
    day_order = {"MONDAY": 1, "TUESDAY": 2, "WEDNESDAY": 3, "THURSDAY": 4, "FRIDAY": 5, "SATURDAY": 6}

    for file_path in files:
        try:
            print(f"\nExtracted from {file_path}:")
            print("=" * 50)
            
            if file_path.endswith('.docx'):
                schedules = process_docx(file_path)
            else:
                schedules = process_pdf(file_path)
                
            if not schedules:
                print("No valid schedule data extracted.")
                continue

            schedules.sort(key=lambda x: (day_order.get(x['day'], 99), x['start']))
            
            for block in schedules:
                print(f"DAY:   {block['day']}")
                print(f"TIME:  {block['start']} to {block['end']}")
                print(f"EVENT: {block['event']}")
                print("-" * 50)
                
        except Exception as e:
            print(f"Failed to process {file_path}: {e}")

if __name__ == "__main__":
    run()